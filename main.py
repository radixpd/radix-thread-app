"""
Analisis Abrasi Benang - RADIX
Tampilan premium dark-navy sesuai desain referensi.
"""

import io
import tempfile
from datetime import datetime

import numpy as np
import pandas as pd
import plotly.graph_objects as go
import streamlit as st
from scipy import interpolate
from sklearn.linear_model import LinearRegression, RANSACRegressor

try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import kaleido  # noqa: F401
    KALEIDO_AVAILABLE = True
except ImportError:
    KALEIDO_AVAILABLE = False

# ─── Page config ────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Analisis Abrasi Benang · RADIX",
    page_icon="🧵",
    layout="wide",
    initial_sidebar_state="collapsed",
)

ACCESS_CODE = "RADIX2025"
TARGET_X_VALUE = 50

INITIAL_DATA = {
    "x_values": [1.7, 3.3, 5.0, 6.7, 8.4, 10.2, 12.0, 13.9, 15.8, 17.7, 19.7, 21.7,
                 23.8, 26.0, 28.2, 30.4, 32.8, 35.3, 37.8, 40.4, 43.3, 46.1, 49.2,
                 52.5, 56.0, 59.9, 64.1, 68.9, 74.66, 82.1],
    "y_values": [105, 143, 157, 185, 191, 191, 200, 250, 266, 292, 337, 343, 345,
                 397, 397, 404, 425, 457, 476, 476, 501, 535, 555, 623, 623, 635,
                 667, 770, 805, 974],
}

# ─── CSS Design System ───────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');

/* ── Reset & Base ── */
*, *::before, *::after { box-sizing: border-box; }

html, body, [class*="css"], .stApp {
    font-family: 'Inter', sans-serif !important;
    background: #EEF2F9 !important;
}

/* Hide Streamlit chrome */
#MainMenu, footer, header, .stDeployButton { visibility: hidden !important; }
.block-container { padding: 0 !important; max-width: 100% !important; }

/* ── HERO BANNER ── */
.hero-banner {
    background: linear-gradient(135deg, #0D1B3E 0%, #1B3A6B 45%, #0D2A55 100%);
    padding: 0;
    position: relative;
    overflow: hidden;
    min-height: 200px;
    display: flex;
    align-items: stretch;
}

.hero-content {
    display: flex;
    align-items: center;
    gap: 28px;
    padding: 32px 48px;
    flex: 1;
    position: relative;
    z-index: 2;
}

.hero-logo-box {
    background: rgba(255,255,255,0.12);
    border: 1.5px solid rgba(255,255,255,0.2);
    border-radius: 20px;
    padding: 18px;
    display: flex;
    align-items: center;
    justify-content: center;
    flex-shrink: 0;
    width: 88px;
    height: 88px;
    font-size: 2.6rem;
    backdrop-filter: blur(10px);
}

.hero-text { flex: 1; }

.hero-eyebrow {
    font-size: 0.72rem;
    font-weight: 600;
    letter-spacing: 4px;
    color: rgba(255,255,255,0.55);
    text-transform: uppercase;
    margin-bottom: 6px;
}

.hero-title {
    font-size: 2.4rem;
    font-weight: 800;
    color: #FFFFFF;
    line-height: 1.1;
    margin: 0 0 8px 0;
}

.hero-subtitle {
    font-size: 0.92rem;
    color: rgba(255,255,255,0.65);
    margin: 0;
}

.hero-subtitle span {
    color: #5B9BFF;
    font-weight: 600;
}

.hero-image-area {
    width: 340px;
    position: relative;
    overflow: hidden;
    flex-shrink: 0;
}

.hero-image-area::before {
    content: '';
    position: absolute;
    inset: 0;
    background: linear-gradient(90deg, #1B3A6B 0%, transparent 40%);
    z-index: 1;
}

.hero-decoration {
    position: absolute;
    right: -20px;
    top: -20px;
    width: 200px;
    height: 200px;
    border-radius: 50%;
    background: radial-gradient(circle, rgba(91,155,255,0.15) 0%, transparent 70%);
}

/* ── NAV BAR ── */
.top-nav {
    background: #0D1B3E;
    padding: 10px 48px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    border-bottom: 1px solid rgba(255,255,255,0.08);
}

.nav-brand {
    display: flex;
    align-items: center;
    gap: 10px;
    color: white;
    font-weight: 700;
    font-size: 0.88rem;
    letter-spacing: 1px;
}

.nav-brand-dot { color: #5B9BFF; }

.nav-tagline {
    color: rgba(255,255,255,0.4);
    font-size: 0.72rem;
    letter-spacing: 0.5px;
}

.nav-badge {
    background: rgba(91,155,255,0.15);
    border: 1px solid rgba(91,155,255,0.3);
    color: #5B9BFF;
    padding: 4px 14px;
    border-radius: 20px;
    font-size: 0.72rem;
    font-weight: 600;
    letter-spacing: 1px;
    text-transform: uppercase;
    cursor: pointer;
    text-decoration: none;
}

/* ── MAIN CONTENT WRAPPER ── */
.content-wrapper {
    padding: 32px 48px;
    max-width: 1400px;
    margin: 0 auto;
}

/* ── SECTION CARDS ── */
.section-card {
    background: #FFFFFF;
    border-radius: 20px;
    padding: 32px;
    margin-bottom: 24px;
    box-shadow: 0 2px 16px rgba(13,27,62,0.07);
    border: 1px solid rgba(13,27,62,0.06);
}

/* ── SECTION HEADER ── */
.section-header {
    display: flex;
    align-items: center;
    gap: 14px;
    margin-bottom: 24px;
}

.section-number {
    width: 36px;
    height: 36px;
    background: linear-gradient(135deg, #1B3A6B, #2D5A9E);
    border-radius: 50%;
    display: flex;
    align-items: center;
    justify-content: center;
    color: white;
    font-weight: 800;
    font-size: 0.9rem;
    flex-shrink: 0;
}

.section-title {
    font-size: 1.25rem;
    font-weight: 700;
    color: #1A2233;
    margin: 0;
}

.section-divider {
    height: 2px;
    background: linear-gradient(90deg, #1B3A6B 0%, transparent 60%);
    border-radius: 2px;
    margin-left: 50px;
    margin-top: -16px;
    margin-bottom: 20px;
    opacity: 0.15;
}

/* ── INFO NOTE ── */
.info-note {
    background: #F0F6FF;
    border: 1px solid #C5D9F7;
    border-radius: 10px;
    padding: 10px 16px;
    font-size: 0.82rem;
    color: #1B3A6B;
    display: flex;
    align-items: center;
    gap: 8px;
}

/* ── TABLE STYLING ── */
.stDataFrame { border-radius: 12px !important; overflow: hidden !important; }
.stDataFrame table { font-size: 0.875rem !important; }
.stDataFrame thead th {
    background: #1B3A6B !important;
    color: white !important;
    font-weight: 600 !important;
    text-align: center !important;
    padding: 12px 16px !important;
}
.stDataFrame tbody tr:hover { background: #F0F6FF !important; }
.stDataFrame tbody td { text-align: center !important; padding: 10px 16px !important; }

/* ── RADIO BUTTONS ── */
.stRadio > div { gap: 12px !important; flex-wrap: wrap !important; }
.stRadio label {
    background: #F5F7FC !important;
    border: 1.5px solid #D0D9EA !important;
    border-radius: 10px !important;
    padding: 8px 18px !important;
    font-size: 0.85rem !important;
    font-weight: 500 !important;
    color: #1A2233 !important;
    cursor: pointer !important;
    transition: all 0.2s !important;
}
.stRadio label:has(input:checked) {
    background: #1B3A6B !important;
    border-color: #1B3A6B !important;
    color: white !important;
}

/* ── METRIC CARDS ── */
.metric-card {
    background: linear-gradient(135deg, #0D1B3E 0%, #1B3A6B 100%);
    border-radius: 16px;
    padding: 24px 28px;
    color: white;
    position: relative;
    overflow: hidden;
    border: 1px solid rgba(255,255,255,0.08);
}

.metric-card::after {
    content: '';
    position: absolute;
    top: -30px;
    right: -30px;
    width: 100px;
    height: 100px;
    border-radius: 50%;
    background: rgba(255,255,255,0.04);
}

.metric-label {
    font-size: 0.75rem;
    font-weight: 600;
    letter-spacing: 2px;
    text-transform: uppercase;
    color: rgba(255,255,255,0.55);
    margin-bottom: 10px;
}

.metric-value {
    font-size: 2rem;
    font-weight: 800;
    color: #FFFFFF;
    line-height: 1;
    margin-bottom: 6px;
}

.metric-caption {
    font-size: 0.75rem;
    color: rgba(255,255,255,0.45);
}

.metric-accent {
    width: 36px;
    height: 3px;
    border-radius: 2px;
    margin-bottom: 14px;
}

/* ── BUTTONS ── */
.stButton > button {
    border-radius: 10px !important;
    font-weight: 600 !important;
    font-size: 0.85rem !important;
    padding: 10px 20px !important;
    transition: all 0.2s !important;
    border: none !important;
}

.stButton > button[kind="primary"] {
    background: linear-gradient(135deg, #1B3A6B, #2D5A9E) !important;
    color: white !important;
}

.stButton > button:not([kind="primary"]) {
    background: #F0F4FA !important;
    color: #1B3A6B !important;
    border: 1.5px solid #D0D9EA !important;
}

.stButton > button:hover {
    transform: translateY(-1px) !important;
    box-shadow: 0 4px 16px rgba(27,58,107,0.25) !important;
}

/* ── DOWNLOAD BUTTON ── */
.stDownloadButton > button {
    background: linear-gradient(135deg, #1B3A6B, #2D5A9E) !important;
    color: white !important;
    border-radius: 10px !important;
    font-weight: 600 !important;
    border: none !important;
    padding: 10px 24px !important;
}

/* ── TABS ── */
.stTabs [data-baseweb="tab-list"] {
    background: #F0F4FA !important;
    border-radius: 12px !important;
    padding: 4px !important;
    gap: 2px !important;
    border: 1px solid #E0E8F5 !important;
}

.stTabs [data-baseweb="tab"] {
    border-radius: 9px !important;
    font-weight: 600 !important;
    font-size: 0.85rem !important;
    color: #6B7A99 !important;
    padding: 8px 20px !important;
}

.stTabs [aria-selected="true"] {
    background: #FFFFFF !important;
    color: #1B3A6B !important;
    box-shadow: 0 1px 6px rgba(0,0,0,0.1) !important;
}

/* ── FILE UPLOADER ── */
.stFileUploader {
    background: #F5F8FF !important;
    border: 2px dashed #C5D9F7 !important;
    border-radius: 12px !important;
    padding: 20px !important;
}

/* ── ALERTS ── */
.stSuccess { border-radius: 10px !important; }
.stError { border-radius: 10px !important; }
.stWarning { border-radius: 10px !important; }
.stInfo { border-radius: 10px !important; }

/* ── DIVIDER ── */
hr { border-color: #E8EEF8 !important; margin: 0 !important; }

/* ── TEXT INPUT ── */
.stTextInput input {
    border-radius: 10px !important;
    border: 1.5px solid #D0D9EA !important;
    font-size: 0.9rem !important;
    padding: 10px 14px !important;
}

.stTextInput input:focus {
    border-color: #1B3A6B !important;
    box-shadow: 0 0 0 3px rgba(27,58,107,0.1) !important;
}

/* ── FOOTER ── */
.footer-bar {
    background: #0D1B3E;
    padding: 20px 48px;
    display: flex;
    align-items: center;
    justify-content: space-between;
    margin-top: 32px;
}

.footer-brand {
    display: flex;
    align-items: center;
    gap: 12px;
}

.footer-logo {
    font-size: 1.5rem;
}

.footer-brand-text { color: white; }
.footer-brand-name {
    font-weight: 700;
    font-size: 0.88rem;
    letter-spacing: 1.5px;
}
.footer-brand-sub {
    font-size: 0.7rem;
    color: rgba(255,255,255,0.4);
    letter-spacing: 0.5px;
}

.footer-tagline {
    font-style: italic;
    color: rgba(255,255,255,0.45);
    font-size: 0.85rem;
}

/* ── PASSWORD GATE ── */
.login-card {
    background: white;
    border-radius: 20px;
    padding: 40px;
    box-shadow: 0 8px 40px rgba(13,27,62,0.15);
    border: 1px solid rgba(13,27,62,0.08);
    max-width: 420px;
    margin: 80px auto;
}

.login-icon {
    font-size: 2.5rem;
    text-align: center;
    margin-bottom: 16px;
}

.login-title {
    font-size: 1.4rem;
    font-weight: 800;
    color: #1A2233;
    text-align: center;
    margin-bottom: 6px;
}

.login-sub {
    font-size: 0.85rem;
    color: #6B7A99;
    text-align: center;
    margin-bottom: 28px;
}

/* Data editor */
[data-testid="stDataEditorDataframe"] {
    border-radius: 12px !important;
}

/* Plotly chart container */
.js-plotly-plot {
    border-radius: 12px;
    overflow: hidden;
}

/* Caption text */
.stCaption { font-size: 0.8rem !important; color: #6B7A99 !important; }

/* Label text inputs */
.stTextInput label { font-weight: 600 !important; color: #1A2233 !important; font-size: 0.85rem !important; }
.stRadio label[data-baseweb] { font-weight: 600 !important; }

/* Column header for data editor */
[data-testid="column-header-0"],
[data-testid="column-header-1"] {
    background: #1B3A6B !important;
    color: white !important;
}

</style>
""", unsafe_allow_html=True)


# ─── Session state ───────────────────────────────────────────────────────────
if "data" not in st.session_state:
    st.session_state.data = pd.DataFrame(INITIAL_DATA)
if "password_entered" not in st.session_state:
    st.session_state.password_entered = False


# ─── Password gate ───────────────────────────────────────────────────────────
def check_password() -> bool:
    if st.session_state.password_entered:
        return True

    st.markdown("""
    <div style="background: linear-gradient(135deg, #0D1B3E 0%, #1B3A6B 100%); 
                min-height: 100vh; display: flex; align-items: center; 
                justify-content: center; padding: 40px;">
    </div>
    """, unsafe_allow_html=True)

    col_a, col_b, col_c = st.columns([1, 1.1, 1])
    with col_b:
        st.markdown("""
        <div class="login-card">
            <div class="login-icon">🔒</div>
            <div class="login-title">Akses Aplikasi</div>
            <div class="login-sub">Masukkan kode akses untuk melanjutkan</div>
        </div>
        """, unsafe_allow_html=True)
        pw = st.text_input("Kode Akses", type="password", placeholder="Masukkan kode akses Anda")
        if st.button("Masuk →", use_container_width=True, type="primary"):
            if pw == ACCESS_CODE:
                st.session_state.password_entered = True
                st.rerun()
            else:
                st.error("Kode akses salah. Silakan coba lagi.")
    return False


if not check_password():
    st.stop()


# ─── Top Navigation Bar ───────────────────────────────────────────────────────
st.markdown("""
<div class="top-nav">
    <div class="nav-brand">
        🧵 &nbsp;
        <div>
            <div>PULCRA <span class="nav-brand-dot">·</span> RADIX</div>
            <div class="nav-tagline">Textile Innovation</div>
        </div>
    </div>
    <a class="nav-badge" href="#">Fork &nbsp; ⑂</a>
</div>
""", unsafe_allow_html=True)


# ─── Hero Banner ─────────────────────────────────────────────────────────────
st.markdown(f"""
<div class="hero-banner">
    <div class="hero-content">
        <div class="hero-logo-box">🧵</div>
        <div class="hero-text">
            <div class="hero-eyebrow">PULCRA · RADIX · TEXTILE INNOVATION</div>
            <h1 class="hero-title">Analisis Abrasi Benang</h1>
            <p class="hero-subtitle">
                Visualisasi data dan perhitungan titik potong pada 
                <span>x = {TARGET_X_VALUE}</span>
            </p>
        </div>
    </div>
    <div class="hero-decoration"></div>
</div>
""", unsafe_allow_html=True)


# ─── Kalkulasi ───────────────────────────────────────────────────────────────
@st.cache_data(show_spinner="Menghitung analisis data…")
def calculate_lines_and_points(x_values_series, y_values_series):
    results = {
        "y_at_x_50_original_curve": np.nan,
        "specific_x1_pt10_20": np.nan, "specific_y1_pt10_20": np.nan,
        "specific_x2_pt10_20": np.nan, "specific_y2_pt10_20": np.nan,
        "y_at_x_50_pt10_20_line": np.nan,
        "pt10_20_line_x_range": np.array([]), "pt10_20_line_y": np.array([]),
        "y_at_x_50_ransac_line": np.nan,
        "ransac_line_x": np.array([]), "ransac_line_y": np.array([]),
    }

    x_np = x_values_series.values
    y_np = y_values_series.values

    if len(x_np) < 2:
        return results

    try:
        if np.all(np.diff(x_np) > 0):
            f = interpolate.interp1d(x_np, y_np, kind="linear", fill_value="extrapolate")
            results["y_at_x_50_original_curve"] = float(f(TARGET_X_VALUE))
    except ValueError:
        pass

    if len(x_np) >= 20:
        results["specific_x1_pt10_20"], results["specific_y1_pt10_20"] = x_np[9], y_np[9]
        results["specific_x2_pt10_20"], results["specific_y2_pt10_20"] = x_np[19], y_np[19]
    elif len(x_np) >= 2:
        results["specific_x1_pt10_20"], results["specific_y1_pt10_20"] = x_np[0], y_np[0]
        results["specific_x2_pt10_20"], results["specific_y2_pt10_20"] = x_np[-1], y_np[-1]

    x1, x2 = results["specific_x1_pt10_20"], results["specific_x2_pt10_20"]
    if not np.isnan(x1) and not np.isnan(x2) and x1 != x2:
        y1, y2 = results["specific_y1_pt10_20"], results["specific_y2_pt10_20"]
        slope = (y2 - y1) / (x2 - x1)
        intercept = y1 - slope * x1
        results["y_at_x_50_pt10_20_line"] = slope * TARGET_X_VALUE + intercept
        x_min, x_max = x_np.min(), x_np.max()
        results["pt10_20_line_x_range"] = np.linspace(min(x_min, TARGET_X_VALUE), max(x_max, TARGET_X_VALUE), 100)
        results["pt10_20_line_y"] = slope * results["pt10_20_line_x_range"] + intercept

    try:
        X_reshaped = x_np.reshape(-1, 1)
        residual_threshold = max(np.std(y_np) * 0.5, 1.0)
        ransac = RANSACRegressor(
            LinearRegression(), min_samples=2, residual_threshold=residual_threshold,
            random_state=42, max_trials=1000,
        )
        ransac.fit(X_reshaped, y_np)
        results["y_at_x_50_ransac_line"] = ransac.predict(np.array([[TARGET_X_VALUE]]))[0]
        x_min, x_max = x_np.min(), x_np.max()
        results["ransac_line_x"] = np.linspace(min(x_min, TARGET_X_VALUE), max(x_max, TARGET_X_VALUE), 100)
        results["ransac_line_y"] = ransac.predict(results["ransac_line_x"].reshape(-1, 1))
    except Exception:
        pass

    return results


# ─── Warna ───────────────────────────────────────────────────────────────────
COLOR_DATA    = "#5B9BFF"
COLOR_TARGET  = "#FF6B6B"
COLOR_PT10_20 = "#F4B740"
COLOR_RANSAC  = "#5FD068"
PLOT_BG       = "#0D1B3E"
GRID_COLOR    = "#1E3060"


def create_abrasion_plot(x_values, y_values, results, analysis_choice):
    fig = go.Figure()
    if x_values.empty:
        return fig

    fig.add_trace(go.Scatter(
        x=x_values, y=y_values, mode="lines+markers", name="Data Aktual",
        line=dict(color=COLOR_DATA, width=2.5),
        marker=dict(size=6, color=COLOR_DATA, line=dict(color="white", width=1)),
    ))

    y_min, y_max = y_values.min(), y_values.max()
    span = max(y_max - y_min, 1)
    y0 = y_min - span * 0.12
    y1 = y_max + span * 0.12

    fig.add_shape(type="line", x0=TARGET_X_VALUE, y0=y0, x1=TARGET_X_VALUE, y1=y1,
                  line=dict(color=COLOR_TARGET, width=1.5, dash="dot"), layer="below")

    fig.add_annotation(
        x=TARGET_X_VALUE, y=y1,
        text=f"<b>x = {TARGET_X_VALUE}</b>",
        showarrow=False, yanchor="bottom",
        font=dict(color="white", size=12, family="Inter"),
        bgcolor=COLOR_TARGET, borderpad=5,
        bordercolor=COLOR_TARGET, borderwidth=1,
    )

    show_pt = analysis_choice in ("Garis Titik 10 & 20", "Tampilkan Semua")
    show_ransac = analysis_choice in ("Garis yang melewati banyak titik", "Tampilkan Semua")
    show_original = analysis_choice in ("Kurva Data Asli", "Tampilkan Semua")

    if show_pt and results.get("pt10_20_line_x_range", np.array([])).size > 0:
        fig.add_trace(go.Scatter(
            x=results["pt10_20_line_x_range"], y=results["pt10_20_line_y"],
            mode="lines", name="Garis Titik 10 & 20",
            line=dict(color=COLOR_PT10_20, width=2, dash="dot"),
        ))
        if not np.isnan(results["y_at_x_50_pt10_20_line"]):
            fig.add_trace(go.Scatter(
                x=[TARGET_X_VALUE], y=[results["y_at_x_50_pt10_20_line"]],
                mode="markers", name="Potongan Garis 10–20",
                marker=dict(size=13, color=COLOR_PT10_20, symbol="star",
                            line=dict(color="white", width=1.5)),
            ))

    if show_ransac and results.get("ransac_line_x", np.array([])).size > 0:
        fig.add_trace(go.Scatter(
            x=results["ransac_line_x"], y=results["ransac_line_y"],
            mode="lines", name="Regresi RANSAC",
            line=dict(color=COLOR_RANSAC, width=2, dash="dash"),
        ))
        if not np.isnan(results["y_at_x_50_ransac_line"]):
            fig.add_trace(go.Scatter(
                x=[TARGET_X_VALUE], y=[results["y_at_x_50_ransac_line"]],
                mode="markers", name="Potongan RANSAC",
                marker=dict(size=13, color=COLOR_RANSAC, symbol="star",
                            line=dict(color="white", width=1.5)),
            ))

    if show_original and not np.isnan(results.get("y_at_x_50_original_curve", np.nan)):
        fig.add_trace(go.Scatter(
            x=[TARGET_X_VALUE], y=[results["y_at_x_50_original_curve"]],
            mode="markers", name="Potongan Kurva Asli",
            marker=dict(size=13, color=COLOR_DATA, symbol="star",
                        line=dict(color="white", width=1.5)),
        ))

    fig.update_layout(
        xaxis_title=dict(text="Nilai X", font=dict(color="rgba(255,255,255,0.7)", size=12)),
        yaxis_title=dict(text="Nilai Benang Putus (N)", font=dict(color="rgba(255,255,255,0.7)", size=12)),
        plot_bgcolor=PLOT_BG,
        paper_bgcolor=PLOT_BG,
        font=dict(color="rgba(255,255,255,0.85)", family="Inter"),
        xaxis=dict(showgrid=True, gridcolor=GRID_COLOR, zeroline=False,
                   tickfont=dict(color="rgba(255,255,255,0.6)"),
                   linecolor=GRID_COLOR),
        yaxis=dict(showgrid=True, gridcolor=GRID_COLOR, zeroline=False,
                   tickfont=dict(color="rgba(255,255,255,0.6)"),
                   linecolor=GRID_COLOR),
        legend=dict(
            orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1,
            bgcolor="rgba(0,0,0,0)", font=dict(size=11),
        ),
        hovermode="x unified",
        margin=dict(l=20, r=20, b=20, t=52),
        height=420,
    )
    return fig


# ─── Content padding wrapper ──────────────────────────────────────────────────
st.markdown('<div style="padding: 32px 48px;">', unsafe_allow_html=True)

# ══════════════════════════════════════════════════════
# SECTION 1 — Input Data
# ══════════════════════════════════════════════════════
st.markdown("""
<div class="section-card">
    <div class="section-header">
        <div class="section-number">1</div>
        <div class="section-title">Input Data</div>
    </div>
</div>
""", unsafe_allow_html=True)

with st.container():
    # Section 1 visual container
    col_hdr, col_note = st.columns([3, 2])
    with col_hdr:
        st.markdown("#### 📋 Input Data")
    with col_note:
        st.markdown("""
        <div class="info-note">
            ℹ️ Nilai X tetap dan tidak dapat diubah. Edit kolom <strong>Nilai Benang Putus (N)</strong> sesuai kebutuhan.
        </div>
        """, unsafe_allow_html=True)

    tab_manual, tab_excel = st.tabs(["✍️  Input Manual", "📁  Impor dari Excel"])

    with tab_manual:
        st.write("")
        edited_data = pd.DataFrame({
            "x_value": st.session_state.data["x_values"],
            "y_value": st.session_state.data["y_values"],
        })
        edited_data.index = edited_data.index + 1

        edited_df = st.data_editor(
            edited_data,
            disabled=["x_value"],
            hide_index=False,
            column_config={
                "x_value": st.column_config.NumberColumn("Nilai Tetap (X)", format="%.1f"),
                "y_value": st.column_config.NumberColumn("Nilai Benang Putus (N)", format="%.2f"),
            },
            num_rows="dynamic",
            use_container_width=True,
            key="data_editor",
        )

        st.write("")
        col1, col2, _ = st.columns([1.2, 1.2, 2])
        with col1:
            if st.button("✅ Terapkan Perubahan", use_container_width=True, type="primary"):
                cleaned = edited_df.dropna(subset=["x_value", "y_value"])
                if cleaned.empty:
                    st.warning("Tabel data kosong.")
                elif not np.all(np.diff(cleaned["x_value"].values) > 0):
                    st.error("Nilai X harus monoton meningkat.")
                elif len(cleaned) != len(INITIAL_DATA["x_values"]):
                    st.warning("Jumlah baris berubah — data dikembalikan ke nilai awal.")
                    st.session_state.data = pd.DataFrame(INITIAL_DATA)
                else:
                    st.session_state.data = pd.DataFrame({
                        "x_values": cleaned["x_value"].values,
                        "y_values": cleaned["y_value"].values,
                    })
                    st.success("Data berhasil diperbarui.")
        with col2:
            if st.button("↺ Reset ke Data Awal", use_container_width=True):
                st.session_state.data = pd.DataFrame(INITIAL_DATA)
                st.success("Data telah direset.")

    with tab_excel:
        st.write("")
        st.caption("Unggah file Excel (.xlsx / .xls) dengan kolom 'x_values' dan 'y_values'.")
        uploaded_file = st.file_uploader("Pilih file Excel", type=["xlsx", "xls"])
        if uploaded_file is not None:
            try:
                df_uploaded = pd.read_excel(uploaded_file)
                if {"x_values", "y_values"}.issubset(df_uploaded.columns):
                    df_uploaded["x_values"] = pd.to_numeric(df_uploaded["x_values"], errors="coerce")
                    df_uploaded["y_values"] = pd.to_numeric(df_uploaded["y_values"], errors="coerce")
                    df_uploaded.dropna(subset=["x_values", "y_values"], inplace=True)
                    if df_uploaded.empty:
                        st.warning("File tidak mengandung data valid.")
                    elif not np.all(np.diff(df_uploaded["x_values"].values) > 0):
                        st.error("Nilai 'x_values' harus monoton meningkat.")
                    else:
                        st.session_state.data = df_uploaded[["x_values", "y_values"]].reset_index(drop=True)
                        st.success("Data berhasil diimpor.")
                        st.dataframe(st.session_state.data.head(), use_container_width=True)
                else:
                    st.error("File Excel harus memiliki kolom 'x_values' dan 'y_values'.")
            except Exception as e:
                st.error(f"Kesalahan saat membaca file: {e}")

st.write("")

# ══════════════════════════════════════════════════════
# SECTION 2 — Analisis & Visualisasi
# ══════════════════════════════════════════════════════
results = calculate_lines_and_points(
    st.session_state.data["x_values"],
    st.session_state.data["y_values"],
)

st.markdown("---")
st.write("")

st.markdown("""
<div style="display:flex; align-items:center; gap:14px; margin-bottom:16px;">
    <div class="section-number">2</div>
    <div class="section-title" style="font-size:1.25rem; font-weight:700; color:#1A2233;">Analisis & Visualisasi</div>
</div>
""", unsafe_allow_html=True)

st.markdown("**Pilih jenis garis analisis:**")
analysis_choice = st.radio(
    "",
    ("Kurva Data Asli", "Garis Titik 10 & 20", "Garis yang melewati banyak titik", "Tampilkan Semua"),
    horizontal=True,
    label_visibility="collapsed",
)

st.write("")

# Chart in dark card
st.markdown("""
<div style="background:#0D1B3E; border-radius:16px; padding:4px; 
            box-shadow:0 4px 24px rgba(13,27,62,0.2); margin-bottom:8px;">
""", unsafe_allow_html=True)

fig = create_abrasion_plot(
    st.session_state.data["x_values"],
    st.session_state.data["y_values"],
    results,
    analysis_choice,
)
st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})

st.markdown("</div>", unsafe_allow_html=True)

st.write("")

# ══════════════════════════════════════════════════════
# SECTION 3 — Hasil Perhitungan
# ══════════════════════════════════════════════════════
st.markdown("---")
st.write("")

st.markdown(f"""
<div style="display:flex; align-items:center; gap:14px; margin-bottom:24px;">
    <div class="section-number">3</div>
    <div class="section-title" style="font-size:1.25rem; font-weight:700; color:#1A2233;">
        Hasil Perhitungan Titik Potong di X = {TARGET_X_VALUE}
    </div>
</div>
""", unsafe_allow_html=True)


def fmt(value) -> str:
    if value is None or (isinstance(value, float) and np.isnan(value)):
        return "—"
    return f"{value:.2f} N"


METRIC_INFO = {
    "Kurva Data Asli": ("y_at_x_50_original_curve", "KURVA DATA ASLI", "#5B9BFF", "Interpolasi linear pada X = 50"),
    "Garis Titik 10 & 20": ("y_at_x_50_pt10_20_line", "GARIS TITIK 10 & 20", "#F4B740", "Garis lurus melalui titik ke-10 dan ke-20"),
    "Garis yang melewati banyak titik": ("y_at_x_50_ransac_line", "GARIS RANSAC", "#5FD068", "Regresi robust, tahan terhadap outlier"),
}

if analysis_choice == "Tampilkan Semua":
    cols = st.columns(3)
    for col, (key, label, accent, caption) in zip(cols, METRIC_INFO.values()):
        val = results.get(key)
        val_str = fmt(val)
        with col:
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-accent" style="background:{accent};"></div>
                <div class="metric-label">{label}</div>
                <div class="metric-value">{val_str}</div>
                <div class="metric-caption">{caption}</div>
            </div>
            """, unsafe_allow_html=True)
else:
    if analysis_choice in METRIC_INFO:
        key, label, accent, caption = METRIC_INFO[analysis_choice]
        val_str = fmt(results.get(key))
        _, mid, _ = st.columns([1, 1, 1])
        with mid:
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-accent" style="background:{accent};"></div>
                <div class="metric-label">{label}</div>
                <div class="metric-value">{val_str}</div>
                <div class="metric-caption">{caption}</div>
            </div>
            """, unsafe_allow_html=True)

st.write("")

# ══════════════════════════════════════════════════════
# SECTION 4 — Unduh Hasil
# ══════════════════════════════════════════════════════
st.markdown("---")
st.write("")

st.markdown("""
<div style="display:flex; align-items:center; gap:14px; margin-bottom:20px;">
    <div class="section-number">4</div>
    <div class="section-title" style="font-size:1.25rem; font-weight:700; color:#1A2233;">
        Unduh Hasil Analisis
    </div>
</div>
""", unsafe_allow_html=True)

if not DOCX_AVAILABLE:
    st.info("Fitur ekspor Word membutuhkan paket `python-docx`. Tambahkan ke requirements.txt untuk mengaktifkan fitur ini.")
else:
    col_input, col_btn, _ = st.columns([2, 1.5, 1.5])
    with col_input:
        filename = st.text_input("Nama file (tanpa ekstensi .docx)", value="Hasil_Analisis_Abrasi",
                                  label_visibility="visible")
    with col_btn:
        st.write("")
        st.write("")
        generate = st.button("📄 Buat Dokumen Word", use_container_width=True, type="primary")

    if generate:
        if not filename.strip():
            st.warning("Masukkan nama file terlebih dahulu.")
        else:
            doc = Document()
            doc.add_heading("Hasil Analisis Abrasi Benang", level=1)
            doc.add_paragraph(f"Dibuat pada: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_heading("Data Abrasi", level=2)
            table = doc.add_table(rows=1, cols=2)
            hdr = table.rows[0].cells
            hdr[0].text, hdr[1].text = "Nilai X", "Nilai Benang Putus (N)"
            for x, y in zip(st.session_state.data["x_values"], st.session_state.data["y_values"]):
                row = table.add_row().cells
                row[0].text, row[1].text = str(x), str(y)

            if KALEIDO_AVAILABLE:
                doc.add_heading("Grafik Analisis", level=2)
                try:
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp_img:
                        fig.write_image(tmp_img.name)
                        doc.add_picture(tmp_img.name, width=Inches(6))
                except Exception as e:
                    doc.add_paragraph(f"(Grafik tidak dapat disertakan: {e})")
            else:
                doc.add_paragraph("(Paket 'kaleido' belum terinstal — grafik tidak disertakan.)")

            doc.add_heading("Hasil Perhitungan", level=2)
            doc.add_paragraph(f"Nilai perpotongan pada X = {TARGET_X_VALUE}:")
            keys_to_show = (
                METRIC_INFO.values() if analysis_choice == "Tampilkan Semua"
                else [METRIC_INFO[analysis_choice]] if analysis_choice in METRIC_INFO else []
            )
            for key, label, _, _ in keys_to_show:
                val = results.get(key)
                if val is not None and not np.isnan(val):
                    doc.add_paragraph(f"{label}: {val:.2f} N", style="List Bullet")

            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            st.success("Dokumen siap diunduh.")
            st.download_button(
                "⬇️ Unduh Dokumen Word",
                data=buf,
                file_name=f"{filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

st.write("")
st.markdown("</div>", unsafe_allow_html=True)

# ─── Footer ──────────────────────────────────────────────────────────────────
st.markdown("""
<div class="footer-bar">
    <div class="footer-brand">
        <div class="footer-logo">🧵</div>
        <div class="footer-brand-text">
            <div class="footer-brand-name">PULCRA · RADIX</div>
            <div class="footer-brand-sub">Innovation in Every Thread</div>
        </div>
    </div>
    <div class="footer-tagline">Stronger Yarn, Better Tomorrow ✨</div>
    <div class="footer-logo">🧶</div>
</div>
""", unsafe_allow_html=True)
